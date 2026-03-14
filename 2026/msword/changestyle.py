# replace_header2_to_normal_bold.py
"""
Replace all Word paragraphs with style "Heading 2" (and common variants) to
"Normal" style with bold formatting in all runs.

Usage:
    python replace_header2_to_normal_bold.py [input_docx] [output_docx]

Defaults:
    input_docx  = mycare.docx
    output_docx = mycare_modified.docx

Requirements:
    - python-docx

Notes:
    - This script preserves existing bold where present and sets bold=True on all runs
      in affected paragraphs.
    - It attempts to be robust to localized or slightly different style names by
      matching case-insensitively against a small set of common aliases.
"""

import sys
from docx import Document

# Common aliases for the Heading 2 style users might refer to
HEADING2_ALIASES = {
    "heading 2",
    "heading2",
    "header2",  # sometimes confused/typed by users
    "h2",
}

def is_heading2_style(paragraph) -> bool:
    """Return True if the paragraph's style matches Heading 2 (by common names)."""
    try:
        name = (paragraph.style.name or "").strip().lower()
    except Exception:
        # Some paragraphs may not have a standard style object; be safe
        name = ""
    if not name:
        return False

    if name in HEADING2_ALIASES:
        return True

    # Also accept names that start with "heading 2" for localized/modified templates
    if name.startswith("heading 2"):
        return True

    # Some templates may embed language tags, e.g., "Heading 2 + Italic"
    if name.replace(" ", "") in ("heading2", "header2"):
        return True

    return False

def convert_heading2_to_normal_bold(doc: Document) -> int:
    """Convert all Heading 2 paragraphs to Normal style with bold runs.

    Returns the number of paragraphs modified.
    """
    modified = 0

    # Iterate through all paragraphs in the document
    for p in doc.paragraphs:
        if is_heading2_style(p):
            # Change paragraph style to Normal
            try:
                p.style = doc.styles["Normal"]
            except KeyError:
                # Fallback: set by name if available or skip if style missing
                p.style = "Normal"

            # Ensure every run is bold
            for run in p.runs:
                run.bold = True

            modified += 1

    # Also traverse paragraphs inside tables (ensure nested tables are handled)
    def iter_table_paragraphs(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for inner_table in cell.tables:
                    yield from iter_table_paragraphs(inner_table)

    for table in doc.tables:
        for p in iter_table_paragraphs(table):
            if is_heading2_style(p):
                try:
                    p.style = doc.styles["Normal"]
                except KeyError:
                    p.style = "Normal"
                for run in p.runs:
                    run.bold = True
                modified += 1

    return modified

def main():
    input_path = sys.argv[1] if len(sys.argv) > 1 else "mycare.docx"
    output_path = sys.argv[2] if len(sys.argv) > 2 else "mycare_modified.docx"

    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Error opening '{input_path}': {e}")
        sys.exit(1)

    modified = convert_heading2_to_normal_bold(doc)

    try:
        doc.save(output_path)
    except Exception as e:
        print(f"Error saving '{output_path}': {e}")
        sys.exit(1)

    print(f"Converted {modified} paragraph(s). Saved to: {output_path}")

if __name__ == "__main__":
    main()